
import os
import docx
import Levenshtein

def getAllDocxFile(argPath):
    retdirs = []
    for (root, dirs, files) in os.walk(argPath):
        print(root, dirs, files)
        for filename in files:
            filePath = os.path.join(root, filename)
            if os.path.splitext(filePath)[1] == '.docx':
                retdirs.append(filePath)

    return retdirs


file = docx.Document('originalDoc.docx')
#print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print(file.paragraphs[i].style.name)
#     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

#print( Levenshtein.ratio(strSource, strStander))

testDir = "/Users/abbylin/百度云同步盘/翻译/永新"
allfiles = getAllDocxFile(testDir)
for file in allfiles:
    print(file)
# if os.path.exists(testDir):
#     print("dir exists")
#     subdirs = os.listdir(testDir)
#     for subdir in subdirs:
#         print(subdir)
#
#     for (root, dirs, files) in os.walk(testDir):
#         for filename in files:
#             filePath = os.path.join(root, filename)
#             if os.path.splitext(filePath)[1] == '.docx':
#                 print(filePath)
#
#             # for dirc in dirs:
#             # print(os.path.join(root, dirc))
# else:
#     print("no dir")


